# backend/reports/exporters/docx_renderer.py

"""
Word (.docx) export renderer for all financial reports.

Uses python-docx. Each report type has its own render function.
All share the same header block, table styling, and page layout.

FEATURES:
    - Company header with report title and date range
    - Professional tables with shaded headers
    - Section headings for multi-section reports
    - Portrait for summary reports, landscape for detail reports
    - Summary paragraph at the bottom

DEPENDENCY:
    pip install python-docx
"""

import io
from decimal import Decimal

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

from .styles import (
    COLOUR_PRIMARY, COLOUR_HEADER_BG, COLOUR_ROW_ALT, COLOUR_BORDER,
    COLOUR_L1, COLOUR_L2, COLOUR_L3, COLOUR_TOTAL, COLOUR_TEXT,
    FONT_NAME, build_header_info, pl_account_amount, pl_l3_amount,
)


# ══════════════════════════════════════════════════
# COLOUR HELPERS
# ══════════════════════════════════════════════════

def _rgb(hex_str):
    """Convert hex colour string to RGBColor."""
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


CLR_PRIMARY = _rgb(COLOUR_PRIMARY)
CLR_HEADER = _rgb(COLOUR_HEADER_BG)
CLR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CLR_TEXT = RGBColor(0x33, 0x33, 0x33)
CLR_GREY = RGBColor(0x77, 0x77, 0x77)

# Classification layer text colours
CLR_L1_TEXT = _rgb(COLOUR_L1)       # Deep indigo
CLR_L2_TEXT = _rgb(COLOUR_L2)       # Strong blue
CLR_L3_TEXT = _rgb(COLOUR_L3)       # Dark teal
CLR_TOTAL_TEXT = _rgb(COLOUR_TOTAL) # Deep indigo
CLR_ACCT_TEXT = _rgb(COLOUR_TEXT)   # Near-black

# Mapping: row style level → text colour
DOCX_ROW_COLOUR_MAP = {
    'L1':    CLR_L1_TEXT,
    'L2':    CLR_L2_TEXT,
    'L3':    CLR_L3_TEXT,
    'TOTAL': CLR_TOTAL_TEXT,
}


# ══════════════════════════════════════════════════
# SHARED DOCX HELPERS
# ══════════════════════════════════════════════════

def _setup_doc(is_landscape=False):
    """Create a new Document with basic settings."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(10)
    font.color.rgb = CLR_TEXT

    # Page orientation
    if is_landscape:
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        # Swap width and height for landscape
        new_width = section.page_height
        new_height = section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    return doc


def _add_header(doc, info, applied_filters=None, extra_meta=None):
    """
    Add company name, report title, date info, optional metadata,
    and an optional "Applied filters:" block.

    Args:
        doc:             python-docx Document instance
        info:            header dict from styles.build_header_info()
        applied_filters: optional list of (label, value) tuples for the
                         "Applied filters:" block. When None or empty,
                         the block is omitted so unfiltered exports stay clean.
        extra_meta:      optional list of (label, value) tuples rendered
                         as a single pipe-separated line directly under the
                         report title — used for totals like "Total entries: 42".

    Keeping `applied_filters` and `extra_meta` as optional kwargs means
    existing renderers that don't pass them behave exactly as before.
    """
    # Company name
    p = doc.add_paragraph()
    run = p.add_run(info['company_name'])
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = CLR_PRIMARY

    # Report title
    p = doc.add_paragraph()
    run = p.add_run(info['report_title'])
    run.font.size = Pt(12)
    run.font.bold = True

    # ── Extra metadata line (e.g., "Total entries: 42") ──
    # Directly under the title, above dates — answers "how big?" before "when?".
    if extra_meta:
        meta_parts = ['{}: {}'.format(label, value) for label, value in extra_meta]
        p = doc.add_paragraph()
        run = p.add_run(' | '.join(meta_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = CLR_GREY

    # Date line
    if info.get('as_of_date'):
        p = doc.add_paragraph()
        run = p.add_run('As of {}'.format(info['as_of_date']))
        run.font.size = Pt(9)
        run.font.color.rgb = CLR_GREY
    elif info.get('from_date') and info.get('to_date'):
        p = doc.add_paragraph()
        run = p.add_run('Period: {} to {}'.format(info['from_date'], info['to_date']))
        run.font.size = Pt(9)
        run.font.color.rgb = CLR_GREY

    # Currency / method line
    parts = []
    if info.get('base_currency'):
        parts.append('Currency: {}'.format(info['base_currency']))
    if info.get('method'):
        parts.append('Method: {}'.format(info['method']))
    if parts:
        p = doc.add_paragraph()
        run = p.add_run(' | '.join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = CLR_GREY

    # ── Applied filters block ──
    # Bold heading, then each filter on its own paragraph with a bold label
    # and plain-text value. Only rendered when the caller provided filters.
    if applied_filters:
        # Small vertical spacer so the block is visually distinct from
        # the currency line above it.
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run('Applied filters:')
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = CLR_TEXT

        for label, value in applied_filters:
            p = doc.add_paragraph()
            # Indent the filter lines slightly — docx has no native CSS
            # left-padding, so we use a leading tab via text content.
            label_run = p.add_run('    {}: '.format(label))
            label_run.font.size = Pt(9)
            label_run.font.bold = True
            label_run.font.color.rgb = CLR_TEXT

            value_run = p.add_run(str(value))
            value_run.font.size = Pt(9)
            value_run.font.color.rgb = CLR_TEXT

    doc.add_paragraph()  # Spacer before the data table


def _add_table(doc, headers, rows, col_widths=None, row_styles=None):
    """
    Add a styled table with thin borders to the document.

    Args:
        doc: Document instance
        headers: list of header strings
        rows: list of lists (each inner list = one row of cell values)
        col_widths: optional list of Inches widths
        row_styles: optional dict {row_index: 'L1'|'L2'|'L3'|'TOTAL'} for
                    bold + coloured classification/total rows.
    """
    row_styles = row_styles or {}
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Apply thin borders to every cell via XML
    WML_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    tbl_pr = table._tbl.tblPr

    for existing in tbl_pr.findall(WML_NS + 'tblBorders'):
        tbl_pr.remove(existing)

    borders_elem = tbl_pr.makeelement(WML_NS + 'tblBorders', {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_elem = borders_elem.makeelement(WML_NS + edge, {})
        edge_elem.set(WML_NS + 'val', 'single')
        edge_elem.set(WML_NS + 'sz', '4')
        edge_elem.set(WML_NS + 'space', '0')
        edge_elem.set(WML_NS + 'color', COLOUR_BORDER)
        borders_elem.append(edge_elem)
    tbl_pr.append(borders_elem)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = CLR_WHITE
        _set_cell_shading(cell, COLOUR_HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        level = row_styles.get(row_idx)  # 'L1', 'L2', 'L3', 'TOTAL', or None
        text_colour = DOCX_ROW_COLOUR_MAP.get(level) if level else None
        is_bold = level is not None

        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(value) if value is not None else ''
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if is_bold:
                        run.font.bold = True
                    if text_colour:
                        run.font.color.rgb = text_colour

        # Alternating row shading
        if row_idx % 2 == 1:
            for cell in row.cells:
                _set_cell_shading(cell, COLOUR_ROW_ALT)

    return table


def _set_cell_shading(cell, colour_hex):
    """Apply a background fill colour to a table cell."""
    WML_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    shading_elem = cell._element.makeelement(WML_NS + 'shd', {})
    shading_elem.set(WML_NS + 'fill', colour_hex)
    shading_elem.set(WML_NS + 'val', 'clear')
    cell._element.get_or_add_tcPr().append(shading_elem)


def _add_summary_line(doc, text, bold=False):
    """Add a summary paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = bold
    return p


def _fmt(value):
    """Format a value for DOCX display."""
    if value is None:
        return ''
    try:
        d = Decimal(str(value))
        return '{:,.2f}'.format(d)
    except Exception:
        return str(value)


def _to_bytes(doc):
    """Save Document to bytes."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════════
# MAIN DISPATCH
# ══════════════════════════════════════════════════

def render_docx(report_type, report_data):
    """Main entry point. Returns bytes (the .docx file content)."""
    renderers = {
        'trial_balance': _render_trial_balance,
        'balance_sheet': _render_balance_sheet,
        'income_statement': _render_income_statement,
        'general_ledger': _render_general_ledger,
        'account_transactions': _render_account_transactions,
        'cash_flow': _render_cash_flow,
        'journal_entries': _render_journal_entries,
    }
    renderer = renderers.get(report_type)
    if not renderer:
        raise ValueError('No DOCX renderer for report type: {}'.format(report_type))
    return renderer(report_data)


# ══════════════════════════════════════════════════
# SHARED RECURSIVE HELPERS
# ══════════════════════════════════════════════════

def _append_acct_rows(rows, acct, depth=3):
    """
    Recursively append account rows to the rows list.
    Shows own balance only (not subtotals) to prevent double-counting.
    Skips zero-balance intermediary accounts but still recurses into children.
    """
    own_dr = acct.get('own_debit_balance')
    own_cr = acct.get('own_credit_balance')
    has_own_balance = (own_dr is not None) or (own_cr is not None)

    if has_own_balance:
        indent = '  ' * depth
        rows.append([
            indent + acct.get('name', ''),
            acct.get('code', ''),
            _fmt(own_dr),
            _fmt(own_cr),
        ])

    for child in acct.get('children', []):
        _append_acct_rows(rows, child, depth + 1)


def _append_section_rows_docx(rows, section, row_styles=None):
    """
    Append L2 > L3 > account rows for a BS section (single Amount column).
    If row_styles dict is provided, L2 and L3 row indices are tracked.
    """
    for l2 in section:
        if row_styles is not None:
            row_styles[len(rows)] = 'L2'
        rows.append(['  ' + l2['name'], '', ''])
        for l3 in l2.get('children', []):
            if row_styles is not None:
                row_styles[len(rows)] = 'L3'
            l3_amt = pl_l3_amount(l3)
            rows.append(['    ' + l3['name'], '', _fmt(l3_amt)])
            for acct in l3.get('accounts', []):
                _append_pl_acct_rows(rows, acct, depth=3)


# ══════════════════════════════════════════════════
# TRIAL BALANCE
# ══════════════════════════════════════════════════

def _render_trial_balance(data):
    doc = _setup_doc()
    info = build_header_info(data)
    _add_header(doc, info)

    rows = []
    row_styles = {}
    for group in data.get('groups', []):
        row_styles[len(rows)] = 'L1'
        rows.append([group['name'], '', '', ''])
        for l2 in group.get('children', []):
            row_styles[len(rows)] = 'L2'
            rows.append(['  ' + l2['name'], '', '', ''])
            for l3 in l2.get('children', []):
                row_styles[len(rows)] = 'L3'
                rows.append(['    ' + l3['name'], '', _fmt(l3.get('subtotal_debit')), _fmt(l3.get('subtotal_credit'))])
                for acct in l3.get('accounts', []):
                    _append_acct_rows(rows, acct, depth=3)

    row_styles[len(rows)] = 'TOTAL'
    rows.append(['TOTAL', '', _fmt(data.get('grand_total_debit')), _fmt(data.get('grand_total_credit'))])

    _add_table(doc, ['Classification / Account', 'Code', 'Debit', 'Credit'], rows, row_styles=row_styles)

    balanced = 'YES' if data.get('is_balanced') else 'NO'
    _add_summary_line(doc, 'Balanced: {}'.format(balanced), bold=True)

    return _to_bytes(doc)


# ══════════════════════════════════════════════════
# BALANCE SHEET (single Amount column)
# ══════════════════════════════════════════════════

def _render_balance_sheet(data):
    doc = _setup_doc()
    info = build_header_info(data)
    _add_header(doc, info)

    rows = []
    row_styles = {}

    # ── ASSETS ──
    row_styles[len(rows)] = 'L1'
    rows.append(['ASSETS', '', ''])
    _append_section_rows_docx(rows, data.get('assets', []), row_styles=row_styles)
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['Total Assets', '', _fmt(data.get('total_assets'))])
    rows.append(['', '', ''])

    # ── LIABILITIES ──
    row_styles[len(rows)] = 'L1'
    rows.append(['LIABILITIES', '', ''])
    _append_section_rows_docx(rows, data.get('liabilities', []), row_styles=row_styles)
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['Total Liabilities', '', _fmt(data.get('total_liabilities'))])
    rows.append(['', '', ''])

    # ── EQUITY ──
    row_styles[len(rows)] = 'L1'
    rows.append(['EQUITY', '', ''])
    _append_section_rows_docx(rows, data.get('equity', []), row_styles=row_styles)
    re_auto = data.get('retained_earnings_auto', {})
    if re_auto:
        rows.append(['  Current Year Net Income (auto)', '', _fmt(re_auto.get('current_year_earnings'))])
        rows.append(['  Prior Year Retained (auto)', '', _fmt(re_auto.get('prior_year_retained'))])
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['Total Equity', '', _fmt(data.get('total_equity'))])
    rows.append(['', '', ''])

    # ── Grand totals ──
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['Total Assets', '', _fmt(data.get('total_assets'))])
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['Total Liabilities & Equity', '', _fmt(data.get('total_liabilities_and_equity'))])

    _add_table(doc, ['Account', 'Code', 'Amount'], rows, row_styles=row_styles)

    balanced = 'YES' if data.get('is_balanced') else 'NO'
    _add_summary_line(doc, 'Balanced: {}'.format(balanced), bold=True)

    return _to_bytes(doc)


# ══════════════════════════════════════════════════
# INCOME STATEMENT (Zoho Books-style P&L — single Amount column)
# ══════════════════════════════════════════════════

def _append_pl_acct_rows(rows, acct, depth=2):
    """
    Recursively append a P&L account row with single Amount.
    Shows own_balance as a positive number; skips zero intermediaries.
    """
    amt = pl_account_amount(acct)
    if amt is not None:
        indent = '  ' * depth
        rows.append([indent + acct.get('name', ''), acct.get('code', ''), _fmt(amt)])
    for child in acct.get('children', []):
        _append_pl_acct_rows(rows, child, depth + 1)


def _append_pl_section_rows_docx(rows, section_data, row_styles=None):
    """
    Append L3 groups + accounts for a P&L section (single Amount column).
    """
    for l2 in section_data:
        for l3 in l2.get('children', []):
            if row_styles is not None:
                row_styles[len(rows)] = 'L3'
            l3_amt = pl_l3_amount(l3)
            rows.append(['  ' + l3['name'], '', _fmt(l3_amt)])
            for acct in l3.get('accounts', []):
                _append_pl_acct_rows(rows, acct, depth=2)


def _render_income_statement(data):
    doc = _setup_doc()
    info = build_header_info(data)
    _add_header(doc, info)

    # Single Amount column
    rows = []
    row_styles = {}

    # ── Operating Income ──
    row_styles[len(rows)] = 'L1'
    rows.append(['OPERATING INCOME', '', ''])
    _append_pl_section_rows_docx(rows, data.get('operating_income', []), row_styles)
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['Total Operating Income', '', _fmt(data.get('total_operating_income'))])
    rows.append(['', '', ''])

    # ── Cost of Goods Sold ──
    row_styles[len(rows)] = 'L1'
    rows.append(['COST OF GOODS SOLD', '', ''])
    _append_pl_section_rows_docx(rows, data.get('cost_of_goods_sold', []), row_styles)
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['Total Cost of Goods Sold', '', _fmt(data.get('total_cogs'))])
    rows.append(['', '', ''])

    # ── Gross Profit ──
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['GROSS PROFIT', '', _fmt(data.get('gross_profit'))])
    rows.append(['', '', ''])

    # ── Operating Expenses ──
    row_styles[len(rows)] = 'L1'
    rows.append(['OPERATING EXPENSES', '', ''])
    _append_pl_section_rows_docx(rows, data.get('operating_expenses', []), row_styles)
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['Total Operating Expenses', '', _fmt(data.get('total_operating_expenses'))])
    rows.append(['', '', ''])

    # ── Operating Profit ──
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['OPERATING PROFIT', '', _fmt(data.get('operating_profit'))])
    rows.append(['', '', ''])

    # ── Non-Operating Income ──
    row_styles[len(rows)] = 'L1'
    rows.append(['NON-OPERATING INCOME', '', ''])
    _append_pl_section_rows_docx(rows, data.get('non_operating_income', []), row_styles)
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['Total Non-Operating Income', '', _fmt(data.get('total_non_operating_income'))])
    rows.append(['', '', ''])

    # ── Non-Operating Expenses ──
    row_styles[len(rows)] = 'L1'
    rows.append(['NON-OPERATING EXPENSES', '', ''])
    _append_pl_section_rows_docx(rows, data.get('non_operating_expenses', []), row_styles)
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['Total Non-Operating Expenses', '', _fmt(data.get('total_non_operating_expenses'))])
    rows.append(['', '', ''])

    # ── Net Profit / Loss ──
    row_styles[len(rows)] = 'TOTAL'
    rows.append(['NET PROFIT / LOSS', '', _fmt(data.get('net_income'))])

    _add_table(doc, ['Account', 'Code', 'Amount'], rows, row_styles=row_styles)

    return _to_bytes(doc)


# ══════════════════════════════════════════════════
# GENERAL LEDGER
# ══════════════════════════════════════════════════

def _render_general_ledger(data):
    doc = _setup_doc(is_landscape=True)
    info = build_header_info(data)
    _add_header(doc, info)

    for account in data.get('accounts', []):
        _add_summary_line(doc, '{} — {} (Opening: {})'.format(
            account['code'], account['name'], _fmt(account.get('opening_balance'))
        ), bold=True)

        rows = []
        for txn in account.get('transactions', []):
            dr_val = _fmt(txn.get('debit')) if txn.get('debit') else ''
            cr_val = _fmt(txn.get('credit')) if txn.get('credit') else ''
            rows.append([
                txn.get('date', ''), dr_val, cr_val,
                _fmt(txn.get('running_balance')),
                str(txn.get('note', ''))[:40], txn.get('journal_type', ''),
            ])
        rows.append(['Closing', _fmt(account.get('total_debit')),
                      _fmt(account.get('total_credit')),
                      _fmt(account.get('closing_balance')), '', ''])

        _add_table(doc, ['Date', 'Debit', 'Credit', 'Running Bal.', 'Note', 'Type'], rows)
        doc.add_paragraph()

    _add_summary_line(doc, 'Grand Total — Debit: {} | Credit: {}'.format(
        _fmt(data.get('grand_total_debit')), _fmt(data.get('grand_total_credit'))
    ), bold=True)

    return _to_bytes(doc)


# ══════════════════════════════════════════════════
# ACCOUNT TRANSACTIONS
# ══════════════════════════════════════════════════

def _render_account_transactions(data):
    doc = _setup_doc(is_landscape=True)
    info = build_header_info(data)
    _add_header(doc, info)

    acct = data.get('account', {})
    _add_summary_line(doc, 'Account: {} — {} ({})'.format(
        acct.get('code', ''), acct.get('name', ''), acct.get('normal_balance', '')
    ), bold=True)
    _add_summary_line(doc, 'Opening Balance: {}'.format(_fmt(data.get('opening_balance'))))

    rows = []
    for txn in data.get('transactions', []):
        dr_val = _fmt(txn.get('debit')) if txn.get('debit') else ''
        cr_val = _fmt(txn.get('credit')) if txn.get('credit') else ''
        rows.append([
            txn.get('date', ''), txn.get('source_number', ''),
            str(txn.get('source_description', ''))[:35],
            dr_val, cr_val,
            _fmt(txn.get('running_balance')),
        ])
    rows.append(['', '', 'CLOSING BALANCE',
                 _fmt(data.get('total_debit')), _fmt(data.get('total_credit')),
                 _fmt(data.get('closing_balance'))])

    _add_table(doc, ['Date', 'Source #', 'Description', 'Debit', 'Credit', 'Running Bal.'], rows)

    return _to_bytes(doc)


# ══════════════════════════════════════════════════
# CASH FLOW STATEMENT
# ══════════════════════════════════════════════════

def _render_cash_flow(data):
    doc = _setup_doc()
    info = build_header_info(data)
    _add_header(doc, info)

    rows = []

    # Operating
    oa = data.get('operating_activities', {})
    rows.append(['OPERATING ACTIVITIES', ''])
    rows.append(['  Net Income', _fmt(oa.get('net_income'))])
    adj = oa.get('adjustments', {})
    rows.append(['  Depreciation & Amortisation (add-back)', _fmt(adj.get('depreciation_and_amortization'))])
    rows.append(['  Working Capital Changes:', ''])
    for item in oa.get('working_capital_changes', []):
        rows.append(['    {} ({})'.format(item['name'], item['code']), _fmt(item.get('cash_effect'))])
    rows.append(['  Net Cash from Operating', _fmt(oa.get('net_cash_from_operating'))])
    rows.append(['', ''])

    # Investing
    ia = data.get('investing_activities', {})
    rows.append(['INVESTING ACTIVITIES', ''])
    for item in ia.get('items', []):
        rows.append(['  {} ({})'.format(item['name'], item['code']), _fmt(item.get('cash_effect'))])
    rows.append(['  Net Cash from Investing', _fmt(ia.get('net_cash_from_investing'))])
    rows.append(['', ''])

    # Financing
    fa = data.get('financing_activities', {})
    rows.append(['FINANCING ACTIVITIES', ''])
    for item in fa.get('items', []):
        rows.append(['  {} ({})'.format(item['name'], item['code']), _fmt(item.get('cash_effect'))])
    rows.append(['  Net Cash from Financing', _fmt(fa.get('net_cash_from_financing'))])
    rows.append(['', ''])

    # Reconciliation
    cr = data.get('cash_reconciliation', {})
    rows.append(['CASH RECONCILIATION', ''])
    rows.append(['  Opening Cash Balance', _fmt(cr.get('opening_cash_balance'))])
    rows.append(['  Net Change in Cash', _fmt(cr.get('net_change_in_cash'))])
    rows.append(['  Closing Cash Balance', _fmt(cr.get('closing_cash_balance'))])

    _add_table(doc, ['Item', 'Amount'], rows)

    balanced = 'YES' if cr.get('is_balanced') else 'NO'
    _add_summary_line(doc, 'Balanced: {}'.format(balanced), bold=True)

    return _to_bytes(doc)


# ══════════════════════════════════════════════════
# JOURNAL ENTRIES
# ══════════════════════════════════════════════════

def _render_journal_entries(data):
    """
    Render Journal Entries export as a landscape DOCX.

    Changes vs. the original implementation:

    1. DESCRIPTION NO LONGER TRUNCATED:
       Word wraps cell content natively — no Paragraph-flowable trick is
       needed (unlike the PDF renderer). The old 30-char truncation dropped
       data for no good reason, so we now pass the full description string.

    2. APPLIED-FILTERS HEADER:
       The header block lists which filters the user applied (status,
       date range, etc.) so anyone reading the DOCX later understands
       what the file contains without guessing from the filename.
    """
    doc = _setup_doc(is_landscape=True)
    info = build_header_info(data)

    # Pass filters and total count through to the extended _add_header.
    # Both are harmless when absent — _add_header silently omits the
    # corresponding sections, so unfiltered exports look identical to before.
    applied_filters = data.get('applied_filters') or []
    journal_count = data.get('journal_count')
    extra_meta = (
        [('Total entries', journal_count)]
        if journal_count is not None
        else None
    )
    _add_header(doc, info, applied_filters=applied_filters, extra_meta=extra_meta)

    rows = []
    for journal in data.get('journals', []):
        for line in journal.get('lines', []):
            dr_val = _fmt(line.get('amount')) if line.get('entry_type') == 'DEBIT' else ''
            cr_val = _fmt(line.get('amount')) if line.get('entry_type') == 'CREDIT' else ''
            rows.append([
                journal.get('entry_number', ''),
                journal.get('date', ''),
                journal.get('status', ''),
                '{} — {}'.format(line.get('account_code', ''), line.get('account_name', '')),
                dr_val,
                cr_val,
                # Full description — Word wraps cell text automatically, so
                # long narratives simply grow the row height. No truncation.
                str(journal.get('description', '')),
            ])

    _add_table(
        doc,
        ['Entry #', 'Date', 'Status', 'Account', 'Debit', 'Credit', 'Description'],
        rows,
    )

    return _to_bytes(doc)